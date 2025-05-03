import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx2pdf import convert
import os
import tkinter as tk
from tkinter import ttk
from datetime import datetime
from config import app_config, ConfigHandler
from db import ClientDB, InvoiceDB, generate_id
from .client_manager import ClientManager
from .invoice_viewer import InvoiceViewer
from .settings_window import SettingsWindow
from .theme import setup_theme

class InvoiceApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Invoice Maker")
        self.state('zoomed') # Full screen
        self.bind('<Escape>', lambda e: self.destroy())  # Press ESC to exit
        self.theme = setup_theme(app_config)
        self.theme_mode = ctk.StringVar(value=app_config['theme_mode'])
        
        # Set window icon
        icon_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'icon.ico')
        if os.path.exists(icon_path):
            self.iconbitmap(icon_path)

        # Initialize variables
        self.current_client_id = ctk.StringVar()
        self.service_totals = [ctk.StringVar(value="0.00") for _ in range(6)]
        self.client_vars = {
            "name": ctk.StringVar(),
            "email": ctk.StringVar(),
            "phone": ctk.StringVar(),
            "address": ctk.StringVar()
        }
        
        # Service variables with trace setup
        self.services = []
        for i in range(6):
            service_vars = {
                "desc": ctk.StringVar(),
                "qty": ctk.StringVar(),
                "price": ctk.StringVar()
            }
            # Add trace to each variable
            for var in service_vars.values():
                var.trace_add("write", lambda *args, idx=i: self.update_service(idx))
            self.services.append(service_vars)
        
        self.tax_percent = ctk.StringVar(value="21")
        self.payment_vars = {
            "method": ctk.StringVar(),
            "entity": ctk.StringVar(),
            "name": ctk.StringVar(),
            "number": ctk.StringVar()
        }
        
        self.load_config()
        self.create_widgets()
        self.load_clients_combobox()

    def load_config(self):
        self.config_data = ConfigHandler.load_config()
        theme = self.config_data.get("theme_mode", "light")
        if theme in self.config_data["color_scheme"]:
            self.color_scheme = self.config_data["color_scheme"][theme]
        else:
            # Fallback to light theme if the requested theme is not available
            self.color_scheme = self.config_data["color_scheme"]["light"]
        self.business_info = self.config_data["business_info"]



    def create_widgets(self):
        main_container = ctk.CTkFrame(self)
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        header_frame = ctk.CTkFrame(main_container)
        header_frame.pack(fill=tk.X, pady=(0, 30))
        
        ctk.CTkLabel(header_frame, text="Invoice Generator", 
                 font=('Inter', 24, 'bold'),
                 text_color=self.color_scheme["primary"]).pack(side=tk.LEFT)
        
        business_info = "\n".join(self.business_info.values())
        ctk.CTkLabel(header_frame, text=business_info, 
                 font=('Inter', 14),
                 justify=tk.RIGHT).pack(side=tk.RIGHT)

        content_frame = ctk.CTkFrame(main_container)
        content_frame.pack(fill=tk.BOTH, expand=True)

        left_panel = ctk.CTkFrame(content_frame, width=400)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 30))

        client_sel_frame = ctk.CTkFrame(left_panel)
        client_sel_frame.pack(fill=tk.X, pady=5)

        ctk.CTkLabel(client_sel_frame, text="Select Client:", font=('Inter', 14)).pack(side=tk.LEFT)
        self.client_cb = ctk.CTkOptionMenu(client_sel_frame, variable=self.current_client_id, 
                                    state="readonly", width=150)
        self.client_cb.pack(side=tk.LEFT, padx=5)
        self.client_cb.configure(command=self.on_client_select)
        
        btn_frame = ctk.CTkFrame(client_sel_frame)
        btn_frame.pack(side=tk.LEFT, padx=10)
        ctk.CTkButton(btn_frame, text="New", command=self.new_client, width=60, font=('Inter', 14)).pack(pady=2)
        ctk.CTkButton(btn_frame, text="Edit", command=self.edit_client, width=60, font=('Inter', 14)).pack(pady=2)

        client_details_frame = ctk.CTkFrame(left_panel)
        client_details_frame.pack(fill=tk.X, pady=10)
        
        ctk.CTkLabel(client_details_frame, text="Client Details", font=('Inter', 16, 'bold')).pack(pady=5)

        fields = [("Name", "name"), ("Email", "email"), ("Phone", "phone"), ("Address", "address")]
        for i, (label, field) in enumerate(fields):
            row_frame = ctk.CTkFrame(client_details_frame)
            row_frame.pack(fill=tk.X, pady=3, padx=10)
            ctk.CTkLabel(row_frame, text=f"{label}:", width=80, anchor=tk.E, font=('Inter', 14)).pack(side=tk.LEFT)
            ctk.CTkEntry(row_frame, textvariable=self.client_vars[field]).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        right_panel = ctk.CTkFrame(content_frame)
        right_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        services_frame = ctk.CTkFrame(right_panel)
        services_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        ctk.CTkLabel(services_frame, text="Services", font=('Inter', 16, 'bold')).pack(pady=5)

        headers = ["Description", "Qty", "Price", "Total"]
        header_frame = ctk.CTkFrame(services_frame)
        header_frame.pack(fill=tk.X, padx=5)
        for col, header in enumerate(headers):
            ctk.CTkLabel(header_frame, text=header, font=('Inter', 14, 'bold'),
                     fg_color=self.color_scheme["primary"],
                     text_color="white",
                     width=120 if col > 0 else 200).pack(side=tk.LEFT, padx=2, pady=5)

        self.service_rows = []
        for i in range(6):
            row_frame = ctk.CTkFrame(services_frame)
            row_frame.pack(fill=tk.X, padx=5, pady=2)
            row = []
            
            for col in range(3):
                state = "normal" if i == 0 else "disabled"
                width = 200 if col == 0 else 120
                entry = ctk.CTkEntry(row_frame, 
                                 textvariable=self.services[i][["desc", "qty", "price"][col]],
                                 state=state,
                                 width=width)
                entry.pack(side=tk.LEFT, padx=2)
                row.append(entry)
            
            total_label = ctk.CTkLabel(row_frame, textvariable=self.service_totals[i],
                                   fg_color=self.color_scheme["secondary"],
                                   text_color="white" if self.theme_mode.get() == "dark" else "black",
                                   width=120)
            total_label.pack(side=tk.LEFT, padx=2)
            self.service_rows.append(row)

        bottom_panel = ctk.CTkFrame(right_panel)
        bottom_panel.pack(fill=tk.X, pady=20)

        tax_frame = ctk.CTkFrame(bottom_panel)
        tax_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        ctk.CTkLabel(tax_frame, text="Tax Settings", font=('Inter', 16, 'bold')).pack(pady=5)
        tax_content = ctk.CTkFrame(tax_frame)
        tax_content.pack(fill=tk.X, padx=10, pady=5)
        ctk.CTkLabel(tax_content, text="Tax Percentage (%):", font=('Inter', 14)).pack(side=tk.LEFT)
        ctk.CTkEntry(tax_content, textvariable=self.tax_percent, width=80).pack(side=tk.LEFT, padx=5)

        payment_frame = ctk.CTkFrame(bottom_panel)
        payment_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        ctk.CTkLabel(payment_frame, text="Payment Details", font=('Inter', 16, 'bold')).pack(pady=5)
        payment_fields = [("Method", "method"), ("Entity", "entity"), 
                        ("Account Name", "name"), ("Account Number", "number")]
        for i, (label, field) in enumerate(payment_fields):
            row_frame = ctk.CTkFrame(payment_frame)
            row_frame.pack(fill=tk.X, pady=2, padx=10)
            ctk.CTkLabel(row_frame, text=f"{label}:", width=100, font=('Inter', 14)).pack(side=tk.LEFT)
            ctk.CTkEntry(row_frame, textvariable=self.payment_vars[field]).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        action_frame = ctk.CTkFrame(bottom_panel)
        action_frame.pack(side=tk.LEFT, padx=20)
        ctk.CTkButton(action_frame, text="Generate Invoice", 
                  command=self.generate_invoice, width=160, font=('Inter', 14)).pack(pady=5)
        ctk.CTkButton(action_frame, text="View Database", 
                  command=self.open_viewer, width=160, font=('Inter', 14)).pack(pady=5)

        footer_frame = ctk.CTkFrame(main_container)
        footer_frame.pack(fill=tk.X, pady=(10, 0))
        
        ctk.CTkButton(footer_frame, text="⚙", 
                  command=self.open_settings, width=100, font=('Inter', 14)).pack(side=tk.LEFT, padx=5)
        
        ctk.CTkButton(footer_frame, text="Toggle Theme", command=self.toggle_theme, width=100, font=('Inter', 14)).pack(side=tk.LEFT, padx=5)

        services_frame.columnconfigure((0,1,2,3), weight=1)
        for i in range(7):
            services_frame.rowconfigure(i, weight=1 if i > 0 else 0)

    def load_clients_combobox(self):
        self.clients = ClientDB.load_clients()
        if self.clients:
            # Sort clients by name for better organization
            self.clients.sort(key=lambda x: x['name'].lower())
            client_options = [f"{c['name']} ({c['id']})" for c in self.clients]
            self.client_cb.configure(values=client_options)
            # Reset selection
            self.client_cb.set("Select a client...")
            # Clear client details
            for field in ["name", "email", "phone", "address"]:
                self.client_vars[field].set("")
        else:
            self.client_cb.configure(values=[])
            self.client_cb.set("No clients available...")
        self.client_cb.configure(dropdown_text_color=self.color_scheme.get("text", "#000000"))

    def on_client_select(self, event):
        selected = self.client_cb.get()
        client_id = selected.split("(")[-1].strip(")")
        client = next((c for c in self.clients if c["id"] == client_id), None)
        if client:
            for field in ["name", "email", "phone", "address"]:
                self.client_vars[field].set(client.get(field, ""))

    def new_client(self):
        dialog = ClientManager(self)
        self.wait_window(dialog)
        if dialog.result:
            client_id = ClientDB.add_client(dialog.result)
            self.load_clients_combobox()
            self.current_client_id.set(f"{dialog.result['name']} ({client_id})")

    def edit_client(self):
        selected = self.client_cb.get()
        if not selected: return
        client_id = selected.split("(")[-1].strip(")")
        client = next((c for c in self.clients if c["id"] == client_id), None)
        if client:
            dialog = ClientManager(self, client)
            self.wait_window(dialog)
            if dialog.result:
                ClientDB.update_client(client_id, dialog.result)
                self.load_clients_combobox()
                self.on_client_select(None)

    def update_service(self, idx):
        try:
            qty = float(self.services[idx]["qty"].get() or 0)
            price = float(self.services[idx]["price"].get() or 0)
            total = qty * price
            self.service_totals[idx].set(f"{total:.2f}")
        except ValueError:
            self.service_totals[idx].set("0.00")
        
        # Enable next row if current has content
        if idx < 5 and any(var.get() for var in self.services[idx].values()):
            next_row = idx + 1
            for widget in self.service_rows[next_row]:
                widget.configure(state="normal")

    def generate_invoice(self):
        placeholders = {
            "[invoice_id]": generate_id(),
            "[date_time]": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "[client_name]": self.client_vars["name"].get(),
            "[client_email]": self.client_vars["email"].get(),
            "[client_phone]": self.client_vars["phone"].get(),
            "[client_adress]": self.client_vars["address"].get(),
            "[business_name]": self.business_info["name"],
            "[business_email]": self.business_info["email"],
            "[business_phone]": self.business_info["phone"],
            "[business_adress]": self.business_info["address"],
            "[tax_%]": self.tax_percent.get(),
            "[payment_method]": self.payment_vars["method"].get(),
            "[payment_entity]": self.payment_vars["entity"].get(),
            "[payment_name]": self.payment_vars["name"].get(),
            "[payment_number]": self.payment_vars["number"].get(),
        }

        subtotal = sum(float(self.services[i]["qty"].get() or 0) * 
                     float(self.services[i]["price"].get() or 0) for i in range(6))

        for i in range(6):
            qty = float(self.services[i]["qty"].get() or 0)
            price = float(self.services[i]["price"].get() or 0)
            placeholders[f"[service{i+1}]"] = self.services[i]["desc"].get()
            placeholders[f"[s{i+1}num]"] = f"{qty:.2f}"
            placeholders[f"[s{i+1}pri]"] = f"{price:.2f}"
            placeholders[f"[s{i+1}sum]"] = f"{qty*price:.2f}"

        tax_percent = float(self.tax_percent.get() or 0)
        iva = subtotal * (tax_percent / 100)
        placeholders["[iva]"] = f"{iva:.2f}"
        placeholders["[total_iva]"] = f"{subtotal + iva:.2f}"

        try:
            doc = Document(self.config_data["template_path"])
        except FileNotFoundError:
            messagebox.showerror("Error", f"Template file not found at {self.config_data['template_path']}")
            return

        for p in doc.paragraphs:
            for key, value in placeholders.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in placeholders.items():
                            if key in p.text:
                                p.text = p.text.replace(key, value)

        temp_doc = "temp_invoice.docx"
        doc.save(temp_doc)

        output_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
            title="Save Invoice As"
        )

        if output_path:
            try:
                convert(temp_doc, output_path)
                invoice_data = {
                    'invoice_id': placeholders['[invoice_id]'],
                    'client_name': placeholders['[client_name]'],
                    'client_email': placeholders['[client_email]'],
                    'client_phone': placeholders['[client_phone]'],
                    'client_address': placeholders['[client_adress]'],
                    'total_amount': float(placeholders['[total_iva]']),
                    'tax_amount': float(placeholders['[iva]']),
                    'invoice_date': placeholders['[date_time]'],
                    'payment_method': placeholders['[payment_method]'],
                    'payment_entity': placeholders['[payment_entity]']
                }
                InvoiceDB.save_invoice(invoice_data)
                messagebox.showinfo("Success", f"Invoice saved to:\n{output_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate PDF:\n{str(e)}")
            finally:
                os.remove(temp_doc)

    def open_viewer(self):
        InvoiceViewer(self)

    def open_settings(self):
        settings_window = SettingsWindow(self)
        self.wait_window(settings_window)
        self.load_config()
        self.configure_styles()

    def configure_styles(self):
        # Configure theme colors for the main window
        self.configure(fg_color=self.color_scheme.get("background", "#ffffff"))
        
        # Define a recursive function to update all widgets
        def update_widget_styles(widget):
            if isinstance(widget, ctk.CTkFrame):
                widget.configure(fg_color=self.color_scheme.get("surface", "#f0f0f0"))
            elif isinstance(widget, ctk.CTkButton):
                widget.configure(
                    fg_color=self.color_scheme.get("primary", "#1f538d"),
                    hover_color=self.color_scheme.get("primary_dark", "#14375e"),
                    text_color=self.color_scheme.get("text", "#000000") if self.theme_mode.get() == "dark" else "white"
                )
            elif isinstance(widget, ctk.CTkLabel):
                widget.configure(text_color=self.color_scheme.get("text", "#000000"))
            elif isinstance(widget, ctk.CTkEntry):
                widget.configure(
                    fg_color=self.color_scheme.get("surface", "#f0f0f0"),
                    text_color=self.color_scheme.get("text", "#000000"),
                    border_color=self.color_scheme.get("secondary", "#8D99AE")
                )
            elif isinstance(widget, ctk.CTkOptionMenu):
                widget.configure(
                    fg_color=self.color_scheme.get("primary", "#1f538d"),
                    button_color=self.color_scheme.get("primary_dark", "#14375e"),
                    button_hover_color=self.color_scheme.get("secondary", "#8D99AE"),
                    text_color=self.color_scheme.get("text", "#000000") if self.theme_mode.get() == "dark" else "white",
                    dropdown_fg_color=self.color_scheme.get("surface", "#f0f0f0"),
                    dropdown_text_color=self.color_scheme.get("text", "#000000"),
                    dropdown_hover_color=self.color_scheme.get("secondary", "#8D99AE")
                )
            
            # Recursively update all children
            if hasattr(widget, 'winfo_children'):
                for child in widget.winfo_children():
                    update_widget_styles(child)
        
        # Apply styles to all widgets
        for widget in self.winfo_children():
            update_widget_styles(widget)
            
        # Update client combobox dropdown text color
        if hasattr(self, 'client_cb'):
            self.client_cb.configure(dropdown_text_color=self.color_scheme.get("text", "#000000"))

    def toggle_theme(self):
        # Toggle between light and dark themes
        new_theme = "dark" if self.theme_mode.get() == "light" else "light"
        self.theme_mode.set(new_theme)
        self.apply_theme(new_theme)

    def apply_theme(self, theme):
        # Update color scheme based on theme
        if theme in self.config_data["color_scheme"]:
            # Apply the theme using the theme module
            from .theme import apply_theme as set_theme
            self.color_scheme = set_theme(theme, self.config_data)
            
            # Update the theme mode in app_config
            app_config['theme_mode'] = theme
            
            # Update CustomTkinter appearance mode
            ctk.set_appearance_mode(theme)
            
            # Refresh UI with new theme
            self.configure_styles()
        else:
            print(f"Theme '{theme}' not found in configuration")
            # Fallback to light theme if the requested theme is not available
            self.theme_mode.set("light")
            self.color_scheme = self.config_data["color_scheme"]["light"]
            ctk.set_appearance_mode("light")