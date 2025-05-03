import json
import sqlite3
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
from docx import Document
from docx2pdf import convert
import os
import random
import string
from datetime import datetime

# ============= DEFAULT CONFIGURATION =============
DEFAULT_CONFIG = {
    "template_path": "invoice_template.docx",
    "clients_db": "clients.json",
    "invoices_db": "invoices.db",
    "business_info": {
        "name": "Your Business Name",
        "email": "business@example.com",
        "phone": "+123 456 7890",
        "address": "123 Business St, City, Country"
    },
    "color_scheme": {
        "background": "#F8F9FA",
        "surface": "#FFFFFF",
        "primary": "#2B2D42",
        "secondary": "#8D99AE",
        "text": "#2B2D42",
        "highlight": "#EF233C"
    }
}

CONFIG_FILE = "app_config.json"

class ConfigHandler:
    @staticmethod
    def load_config():
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            with open(CONFIG_FILE, "w") as f:
                json.dump(DEFAULT_CONFIG, f, indent=2)
            return DEFAULT_CONFIG

    @staticmethod
    def save_config(config):
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f, indent=2)

app_config = ConfigHandler.load_config()

# ============= DATABASE HANDLERS =============
class ClientDB:
    @staticmethod
    def load_clients():
        try:
            with open(app_config["clients_db"], "r") as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            return []

    @staticmethod
    def save_clients(clients):
        with open(app_config["clients_db"], "w") as f:
            json.dump(clients, f, indent=2)

    @staticmethod
    def add_client(client_data):
        clients = ClientDB.load_clients()
        client_data["id"] = generate_id("CLT")
        clients.append(client_data)
        ClientDB.save_clients(clients)
        return client_data["id"]

    @staticmethod
    def update_client(client_id, new_data):
        clients = ClientDB.load_clients()
        for client in clients:
            if client["id"] == client_id:
                client.update(new_data)
                break
        ClientDB.save_clients(clients)

class InvoiceDB:
    @staticmethod
    def initialize():
        conn = sqlite3.connect(app_config["invoices_db"])
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS invoices
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      invoice_id TEXT UNIQUE,
                      client_name TEXT,
                      client_email TEXT,
                      client_phone TEXT,
                      client_address TEXT,
                      total_amount REAL,
                      tax_amount REAL,
                      invoice_date TEXT,
                      payment_method TEXT,
                      payment_entity TEXT)''')
        conn.commit()
        conn.close()

    @staticmethod
    def save_invoice(invoice_data):
        conn = sqlite3.connect(app_config["invoices_db"])
        c = conn.cursor()
        c.execute('''INSERT INTO invoices 
                     (invoice_id, client_name, client_email, client_phone, client_address,
                      total_amount, tax_amount, invoice_date, payment_method, payment_entity)
                     VALUES (?,?,?,?,?,?,?,?,?,?)''',
                  (invoice_data['invoice_id'],
                   invoice_data['client_name'],
                   invoice_data['client_email'],
                   invoice_data['client_phone'],
                   invoice_data['client_address'],
                   invoice_data['total_amount'],
                   invoice_data['tax_amount'],
                   invoice_data['invoice_date'],
                   invoice_data['payment_method'],
                   invoice_data['payment_entity']))
        conn.commit()
        conn.close()

    @staticmethod
    def get_all_invoices():
        conn = sqlite3.connect(app_config["invoices_db"])
        c = conn.cursor()
        c.execute('''SELECT invoice_id, client_name, invoice_date, 
                     total_amount, tax_amount, payment_method 
                     FROM invoices ORDER BY invoice_date DESC''')
        invoices = c.fetchall()
        conn.close()
        return invoices

    @staticmethod
    def get_invoice_details(invoice_id):
        conn = sqlite3.connect(app_config["invoices_db"])
        c = conn.cursor()
        c.execute('''SELECT * FROM invoices WHERE invoice_id = ?''', (invoice_id,))
        details = c.fetchone()
        conn.close()
        return details

InvoiceDB.initialize()

# ============= UTILITY FUNCTIONS =============
def generate_id(prefix="INV"):
    date_str = datetime.now().strftime("%y%m%d")
    random_str = ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))
    return f"{prefix}-{date_str}-{random_str}"

# ============= GUI COMPONENTS =============
class ClientManager(tk.Toplevel):
    def __init__(self, parent, client_data=None):
        super().__init__(parent)
        self.title("Manage Client" if client_data else "New Client")
        self.client_data = client_data or {}
        self.result = None
        self.create_widgets()
        self.configure(bg=app_config["color_scheme"]["background"])

    def create_widgets(self):
        main_frame = ttk.Frame(self, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        fields = [
            ("Name:", "name"),
            ("Email:", "email"),
            ("Phone:", "phone"),
            ("Address:", "address")
        ]

        for i, (label, field) in enumerate(fields):
            row_frame = ttk.Frame(main_frame)
            row_frame.pack(fill=tk.X, pady=4)
            ttk.Label(row_frame, text=label, width=10, anchor=tk.E).pack(side=tk.LEFT)
            entry = ttk.Entry(row_frame)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
            if self.client_data:
                entry.insert(0, self.client_data.get(field, ""))
            setattr(self, f"{field}_entry", entry)

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=10)
        ttk.Button(btn_frame, text="Save", command=self.save).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(side=tk.LEFT)

    def save(self):
        data = {
            "name": self.name_entry.get(),
            "email": self.email_entry.get(),
            "phone": self.phone_entry.get(),
            "address": self.address_entry.get()
        }
        if not data["name"]:
            messagebox.showerror("Error", "Name is required")
            return
        self.result = data
        self.destroy()

class InvoiceViewer(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Invoice Database Viewer")
        self.geometry("1200x700")
        self.configure(bg=app_config["color_scheme"]["background"])
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        self.tree = ttk.Treeview(main_frame, columns=('ID', 'Client', 'Date', 'Total', 'Tax', 'Method'), show='headings')
        vsb = ttk.Scrollbar(main_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(main_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        headers = [
            ('ID', 'Invoice ID', 200),
            ('Client', 'Client', 250),
            ('Date', 'Date', 120),
            ('Total', 'Total', 100),
            ('Tax', 'Tax', 100),
            ('Method', 'Payment Method', 150)
        ]
        
        for col, text, width in headers:
            self.tree.heading(col, text=text)
            self.tree.column(col, width=width, anchor=tk.CENTER)

        self.tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        
        self.tree.bind("<Double-1>", self.show_details)
        self.load_invoices()

    def load_invoices(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        for inv in InvoiceDB.get_all_invoices():
            self.tree.insert('', 'end', values=inv)

    def show_details(self, event):
        item = self.tree.selection()[0]
        invoice_id = self.tree.item(item, 'values')[0]
        details = InvoiceDB.get_invoice_details(invoice_id)
        
        detail_window = tk.Toplevel(self)
        detail_window.title(f"Invoice Details - {invoice_id}")
        detail_window.configure(bg=app_config["color_scheme"]["background"])
        
        fields = [
            ("Invoice ID:", details[1]),
            ("Client Name:", details[2]),
            ("Client Email:", details[3]),
            ("Client Phone:", details[4]),
            ("Client Address:", details[5]),
            ("Total Amount:", f"${details[6]:.2f}"),
            ("Tax Amount:", f"${details[7]:.2f}"),
            ("Invoice Date:", details[8]),
            ("Payment Method:", details[9]),
            ("Payment Entity:", details[10])
        ]
        
        for i, (label, value) in enumerate(fields):
            ttk.Label(detail_window, text=label, anchor=tk.E).grid(row=i, column=0, padx=5, pady=2, sticky='e')
            ttk.Label(detail_window, text=value, anchor=tk.W).grid(row=i, column=1, padx=5, pady=2, sticky='w')

class SettingsWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Application Settings")
        self.config = app_config.copy()
        self.configure(bg=app_config["color_scheme"]["background"])
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=tk.BOTH, expand=True)

        general_frame = ttk.Frame(notebook)
        notebook.add(general_frame, text="General")
        self.create_path_settings(general_frame)
        
        biz_frame = ttk.Frame(notebook)
        notebook.add(biz_frame, text="Business")
        self.create_business_settings(biz_frame)
        
        appearance_frame = ttk.Frame(notebook)
        notebook.add(appearance_frame, text="Appearance")
        self.create_appearance_settings(appearance_frame)

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=20)
        ttk.Button(btn_frame, text="Save", command=self.save_settings).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(side=tk.RIGHT)

    def create_path_settings(self, parent):
        frame = ttk.LabelFrame(parent, text="File Paths", padding=15)
        frame.pack(fill=tk.BOTH, pady=5, padx=5)
        self.create_path_row(frame, "Template Path:", "template_path", 0)
        self.create_path_row(frame, "Clients DB:", "clients_db", 1)
        self.create_path_row(frame, "Invoices DB:", "invoices_db", 2)

    def create_business_settings(self, parent):
        frame = ttk.LabelFrame(parent, text="Business Information", padding=15)
        frame.pack(fill=tk.BOTH, pady=5, padx=5)
        biz_fields = [
            ("Name:", "business_info", "name"),
            ("Email:", "business_info", "email"),
            ("Phone:", "business_info", "phone"),
            ("Address:", "business_info", "address")
        ]
        for i, (label, section, key) in enumerate(biz_fields):
            self.create_entry_row(frame, label, section, key, i)

    def create_appearance_settings(self, parent):
        frame = ttk.LabelFrame(parent, text="Color Scheme", padding=15)
        frame.pack(fill=tk.BOTH, pady=5, padx=5)
        color_fields = [
            ("Background:", "color_scheme", "background"),
            ("Surface:", "color_scheme", "surface"),
            ("Primary:", "color_scheme", "primary"),
            ("Secondary:", "color_scheme", "secondary"),
            ("Text:", "color_scheme", "text"),
            ("Highlight:", "color_scheme", "highlight")
        ]
        for i, (label, section, key) in enumerate(color_fields):
            self.create_color_row(frame, label, section, key, i)

    def create_path_row(self, parent, label, key, row):
        frame = ttk.Frame(parent)
        frame.pack(fill=tk.X, pady=4)
        ttk.Label(frame, text=label, width=15).pack(side=tk.LEFT)
        entry = ttk.Entry(frame)
        entry.insert(0, self.config[key])
        entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Button(frame, text="Browse...", 
                 command=lambda k=key, e=entry: self.browse_file(k, e)).pack(side=tk.LEFT)
        setattr(self, f"{key}_entry", entry)

    def create_entry_row(self, parent, label, section, key, row):
        frame = ttk.Frame(parent)
        frame.pack(fill=tk.X, pady=4)
        ttk.Label(frame, text=label, width=15).pack(side=tk.LEFT)
        entry = ttk.Entry(frame)
        entry.insert(0, self.config[section][key])
        entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        setattr(self, f"{section}_{key}_entry", entry)

    def create_color_row(self, parent, label, section, key, row):
        frame = ttk.Frame(parent)
        frame.pack(fill=tk.X, pady=4)
        ttk.Label(frame, text=label, width=15).pack(side=tk.LEFT)
        entry = ttk.Entry(frame, width=10)
        entry.insert(0, self.config[section][key])
        entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(frame, text="Pick", 
                 command=lambda e=entry: self.pick_color(e)).pack(side=tk.LEFT)
        setattr(self, f"{section}_{key}_entry", entry)

    def browse_file(self, key, entry):
        if key == "template_path":
            path = filedialog.askopenfilename(filetypes=[("Word Templates", "*.docx")])
        else:
            path = filedialog.asksaveasfilename()
        if path:
            entry.delete(0, tk.END)
            entry.insert(0, path)

    def pick_color(self, entry):
        color = colorchooser.askcolor(title="Choose color", initialcolor=entry.get())
        if color[1]:
            entry.delete(0, tk.END)
            entry.insert(0, color[1])

    def save_settings(self):
        self.config["template_path"] = self.template_path_entry.get()
        self.config["clients_db"] = self.clients_db_entry.get()
        self.config["invoices_db"] = self.invoices_db_entry.get()

        biz_keys = ["name", "email", "phone", "address"]
        for key in biz_keys:
            self.config["business_info"][key] = getattr(self, f"business_info_{key}_entry").get()

        color_keys = ["background", "surface", "primary", "secondary", "text", "highlight"]
        for key in color_keys:
            self.config["color_scheme"][key] = getattr(self, f"color_scheme_{key}_entry").get()

        ConfigHandler.save_config(self.config)
        messagebox.showinfo("Settings Saved", 
                          "Some changes may require application restart to take effect")
        self.destroy()

class InvoiceApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Invoice Generator")
        self.state('zoomed') # Full screen
        self.bind('<Escape>', lambda e: self.destroy())  # Press ESC to exit

        
        # Initialize variables
        self.current_client_id = tk.StringVar()
        self.service_totals = [tk.StringVar(value="0.00") for _ in range(6)]
        self.client_vars = {
            "name": tk.StringVar(),
            "email": tk.StringVar(),
            "phone": tk.StringVar(),
            "address": tk.StringVar()
        }
        
        # Service variables with trace setup
        self.services = []
        for i in range(6):
            service_vars = {
                "desc": tk.StringVar(),
                "qty": tk.StringVar(),
                "price": tk.StringVar()
            }
            # Add trace to each variable
            for var in service_vars.values():
                var.trace_add("write", lambda *args, idx=i: self.update_service(idx))
            self.services.append(service_vars)
        
        self.tax_percent = tk.StringVar(value="21")
        self.payment_vars = {
            "method": tk.StringVar(),
            "entity": tk.StringVar(),
            "name": tk.StringVar(),
            "number": tk.StringVar()
        }
        
        self.load_config()
        self.configure_styles()
        self.create_widgets()
        self.load_clients_combobox()

    def load_config(self):
        self.config = ConfigHandler.load_config()
        self.color_scheme = self.config["color_scheme"]
        self.business_info = self.config["business_info"]

    def configure_styles(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        self.style.configure(".", 
            font=('Segoe UI', 11), 
            background=self.color_scheme["background"],
            foreground=self.color_scheme["text"]
        )
        
        self.style.configure("TFrame", 
            background=self.color_scheme["background"]
        )
        self.style.configure("TLabel", 
            background=self.color_scheme["background"],
            foreground=self.color_scheme["text"],
            padding=5
        )
        self.style.configure("TButton", 
            background=self.color_scheme["secondary"],
            foreground=self.color_scheme["text"],
            borderwidth=0,
            padding=8,
            focuscolor=self.color_scheme["surface"]
        )
        self.style.configure("TEntry", 
            fieldbackground=self.color_scheme["surface"],
            bordercolor=self.color_scheme["secondary"],
            relief="flat",
            padding=6
        )
        self.style.configure("TCombobox", 
            fieldbackground=self.color_scheme["surface"],
            selectbackground=self.color_scheme["surface"]
        )
        self.style.configure("Treeview", 
            background=self.color_scheme["surface"],
            fieldbackground=self.color_scheme["surface"],
            foreground=self.color_scheme["text"],
            rowheight=28
        )
        self.style.configure("Treeview.Heading", 
            background=self.color_scheme["primary"],
            foreground="white",
            font=('Segoe UI', 10, 'bold')
        )

    def create_widgets(self):
        main_container = ttk.Frame(self)
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        header_frame = ttk.Frame(main_container)
        header_frame.pack(fill=tk.X, pady=(0, 30))
        
        ttk.Label(header_frame, text="Invoice Generator", 
                 font=('Segoe UI', 24, 'bold'),
                 foreground=self.color_scheme["primary"]).pack(side=tk.LEFT)
        
        business_info = "\n".join(self.business_info.values())
        ttk.Label(header_frame, text=business_info, 
                 font=('Segoe UI', 11),
                 justify=tk.RIGHT).pack(side=tk.RIGHT)

        content_frame = ttk.Frame(main_container)
        content_frame.pack(fill=tk.BOTH, expand=True)

        left_panel = ttk.Frame(content_frame, width=400)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 30))

        client_sel_frame = ttk.LabelFrame(left_panel, text="Client Management", padding=15)
        client_sel_frame.pack(fill=tk.X, pady=5)

        ttk.Label(client_sel_frame, text="Select Client:").pack(side=tk.LEFT)
        self.client_cb = ttk.Combobox(client_sel_frame, textvariable=self.current_client_id, 
                                    state="readonly", width=25)
        self.client_cb.pack(side=tk.LEFT, padx=5)
        self.client_cb.bind("<<ComboboxSelected>>", self.on_client_select)
        
        btn_frame = ttk.Frame(client_sel_frame)
        btn_frame.pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="New", command=self.new_client, width=6).pack(pady=2)
        ttk.Button(btn_frame, text="Edit", command=self.edit_client, width=6).pack(pady=2)

        client_details_frame = ttk.LabelFrame(left_panel, text="Client Details", padding=15)
        client_details_frame.pack(fill=tk.X, pady=10)

        fields = [("Name", "name"), ("Email", "email"), ("Phone", "phone"), ("Address", "address")]
        for i, (label, field) in enumerate(fields):
            row_frame = ttk.Frame(client_details_frame)
            row_frame.pack(fill=tk.X, pady=3)
            ttk.Label(row_frame, text=f"{label}:", width=10, anchor=tk.E).pack(side=tk.LEFT)
            ttk.Entry(row_frame, textvariable=self.client_vars[field]).pack(side=tk.LEFT, fill=tk.X, expand=True)

        right_panel = ttk.Frame(content_frame)
        right_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        services_frame = ttk.LabelFrame(right_panel, text="Services", padding=15)
        services_frame.pack(fill=tk.BOTH, expand=True)

        headers = ["Description", "Qty", "Price", "Total"]
        for col, header in enumerate(headers):
            ttk.Label(services_frame, text=header, font=('Segoe UI', 10, 'bold'),
                     background=self.color_scheme["primary"],
                     foreground="white",
                     padding=5).grid(
                     row=0, column=col, sticky="ew", padx=1, pady=1)

        self.service_rows = []
        for i in range(6):
            row = []
            for col in range(3):
                state = "normal" if i == 0 else "disabled"
                entry = ttk.Entry(services_frame, 
                                 textvariable=self.services[i][["desc", "qty", "price"][col]],
                                 state=state)
                entry.grid(row=i+1, column=col, padx=1, pady=1, sticky="ew")
                row.append(entry)
            
            total_label = ttk.Label(services_frame, textvariable=self.service_totals[i],
                                   background=self.color_scheme["surface"],
                                   padding=5)
            total_label.grid(row=i+1, column=3, padx=1, pady=1, sticky="ew")
            self.service_rows.append(row)

        bottom_panel = ttk.Frame(right_panel)
        bottom_panel.pack(fill=tk.X, pady=20)

        tax_frame = ttk.LabelFrame(bottom_panel, text="Tax Settings", padding=10)
        tax_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        ttk.Label(tax_frame, text="Tax Percentage (%):").pack(side=tk.LEFT)
        ttk.Entry(tax_frame, textvariable=self.tax_percent, width=8).pack(side=tk.LEFT, padx=5)

        payment_frame = ttk.LabelFrame(bottom_panel, text="Payment Details", padding=10)
        payment_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        payment_fields = [("Method", "method"), ("Entity", "entity"), 
                        ("Account Name", "name"), ("Account Number", "number")]
        for i, (label, field) in enumerate(payment_fields):
            row_frame = ttk.Frame(payment_frame)
            row_frame.pack(fill=tk.X, pady=2)
            ttk.Label(row_frame, text=f"{label}:", width=12).pack(side=tk.LEFT)
            ttk.Entry(row_frame, textvariable=self.payment_vars[field]).pack(side=tk.LEFT, fill=tk.X, expand=True)

        action_frame = ttk.Frame(bottom_panel)
        action_frame.pack(side=tk.LEFT, padx=20)
        ttk.Button(action_frame, text="Generate Invoice", 
                  command=self.generate_invoice, width=18).pack(pady=5)
        ttk.Button(action_frame, text="View Database", 
                  command=self.open_viewer, width=18).pack(pady=5)

        footer_frame = ttk.Frame(main_container)
        footer_frame.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(footer_frame, text="⚙ Settings", 
                  command=self.open_settings).pack(side=tk.LEFT, padx=5)

        services_frame.columnconfigure((0,1,2,3), weight=1)
        for i in range(7):
            services_frame.rowconfigure(i, weight=1 if i > 0 else 0)

    def load_clients_combobox(self):
        self.clients = ClientDB.load_clients()
        self.client_cb["values"] = [f"{c['name']} ({c['id']})" for c in self.clients]

    def on_client_select(self, event):
        selected = self.client_cb.get()
        client_id = selected.split("(")[-1].strip(")")
        client = next((c for c in self.clients if c["id"] == client_id), None)
        if client:
            for field in ["name", "email", "phone", "address"]:
                self.client_vars[field].set(client.get(field, ""))

    def new_client(self):
        dialog = ClientManager(self)
        self.wait_window(dialog)
        if dialog.result:
            client_id = ClientDB.add_client(dialog.result)
            self.load_clients_combobox()
            self.current_client_id.set(f"{dialog.result['name']} ({client_id})")

    def edit_client(self):
        selected = self.client_cb.get()
        if not selected: return
        client_id = selected.split("(")[-1].strip(")")
        client = next((c for c in self.clients if c["id"] == client_id), None)
        if client:
            dialog = ClientManager(self, client)
            self.wait_window(dialog)
            if dialog.result:
                ClientDB.update_client(client_id, dialog.result)
                self.load_clients_combobox()
                self.on_client_select(None)

    def update_service(self, idx):
        try:
            qty = float(self.services[idx]["qty"].get() or 0)
            price = float(self.services[idx]["price"].get() or 0)
            total = qty * price
            self.service_totals[idx].set(f"{total:.2f}")
        except ValueError:
            self.service_totals[idx].set("0.00")
        
        # Enable next row if current has content
        if idx < 5 and any(var.get() for var in self.services[idx].values()):
            next_row = idx + 1
            for widget in self.service_rows[next_row]:
                widget.config(state="normal")

    def generate_invoice(self):
        placeholders = {
            "[invoice_id]": generate_id(),
            "[date_time]": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "[client_name]": self.client_vars["name"].get(),
            "[client_email]": self.client_vars["email"].get(),
            "[client_phone]": self.client_vars["phone"].get(),
            "[client_adress]": self.client_vars["address"].get(),
            "[business_name]": self.business_info["name"],
            "[business_email]": self.business_info["email"],
            "[business_phone]": self.business_info["phone"],
            "[business_adress]": self.business_info["address"],
            "[tax_%]": self.tax_percent.get(),
            "[payment_method]": self.payment_vars["method"].get(),
            "[payment_entity]": self.payment_vars["entity"].get(),
            "[payment_name]": self.payment_vars["name"].get(),
            "[payment_number]": self.payment_vars["number"].get(),
        }

        subtotal = sum(float(self.services[i]["qty"].get() or 0) * 
                     float(self.services[i]["price"].get() or 0) for i in range(6))

        for i in range(6):
            qty = float(self.services[i]["qty"].get() or 0)
            price = float(self.services[i]["price"].get() or 0)
            placeholders[f"[service{i+1}]"] = self.services[i]["desc"].get()
            placeholders[f"[s{i+1}num]"] = f"{qty:.2f}"
            placeholders[f"[s{i+1}pri]"] = f"{price:.2f}"
            placeholders[f"[s{i+1}sum]"] = f"{qty*price:.2f}"

        tax_percent = float(self.tax_percent.get() or 0)
        iva = subtotal * (tax_percent / 100)
        placeholders["[iva]"] = f"{iva:.2f}"
        placeholders["[total_iva]"] = f"{subtotal + iva:.2f}"

        try:
            doc = Document(self.config["template_path"])
        except FileNotFoundError:
            messagebox.showerror("Error", f"Template file not found at {self.config['template_path']}")
            return

        for p in doc.paragraphs:
            for key, value in placeholders.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in placeholders.items():
                            if key in p.text:
                                p.text = p.text.replace(key, value)

        temp_doc = "temp_invoice.docx"
        doc.save(temp_doc)

        output_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
            title="Save Invoice As"
        )

        if output_path:
            try:
                convert(temp_doc, output_path)
                invoice_data = {
                    'invoice_id': placeholders['[invoice_id]'],
                    'client_name': placeholders['[client_name]'],
                    'client_email': placeholders['[client_email]'],
                    'client_phone': placeholders['[client_phone]'],
                    'client_address': placeholders['[client_adress]'],
                    'total_amount': float(placeholders['[total_iva]']),
                    'tax_amount': float(placeholders['[iva]']),
                    'invoice_date': placeholders['[date_time]'],
                    'payment_method': placeholders['[payment_method]'],
                    'payment_entity': placeholders['[payment_entity]']
                }
                InvoiceDB.save_invoice(invoice_data)
                messagebox.showinfo("Success", f"Invoice saved to:\n{output_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate PDF:\n{str(e)}")
            finally:
                os.remove(temp_doc)

    def open_viewer(self):
        InvoiceViewer(self)

    def open_settings(self):
        SettingsWindow(self)
        self.wait_window(SettingsWindow(self))
        self.load_config()
        self.configure_styles()

if __name__ == "__main__":
    app = InvoiceApp()
    
    # Construir la ruta al icono relativa al script
    script_dir = os.path.dirname(__file__) # Directorio del script
    icon_path = os.path.join(script_dir, "icon.ico") 

    try:
        if os.path.exists(icon_path):
             app.iconbitmap(icon_path) # Usar la ruta completa/relativa segura
        else:
             print("Advertencia: Archivo de icono 'icon.ico' no encontrado.")
    except tk.TclError as e:
        print(f"Advertencia: No se pudo cargar el icono '{icon_path}'. Error: {e}")
        
    app.mainloop()
    